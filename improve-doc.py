import openai
import os
import argparse
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert


parser = argparse.ArgumentParser()
parser.add_argument('original_documentFile')
parser.add_argument('new_documentFile')
parser.add_argument('language_summarize')
args = parser.parse_args()

if str(args.language_summarize).lower() == "en":
    conditional = "Summarize the paragraph"
else:
    conditional = "Resuma o paragrafo"

def summarize_paragraph(paragraph):
    openai.api_key = os.environ["open_key"]

    prompt = f"{conditional}: {paragraph}"
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=200,
        temperature=0.5,
        top_p=1.0,
        frequency_penalty=0.0,
        presence_penalty=0.0
    )

    return response.choices[0].text.strip()

def extract_paragraphs(original_file_name):
    original_document = Document(original_file_name)
    paragraphs = [paragraph.text for paragraph in original_document.paragraphs]
    images = [run for paragraph in original_document.paragraphs for run in paragraph.runs if run.element.tag.endswith(('}drawing', '}pict'))]
    tables = original_document.tables

    return paragraphs, images, tables

def copy_tables(original_tables, output_document):
    for table in original_tables:
        new_table = output_document.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = cell.text

    return output_document

def copy_images(original_images, output_document):
    for image in original_images:
        image_width = image.element.attrib.get('cx')
        image_height = image.element.attrib.get('cy')
        image_id = image._r.get_or_add_drawing().get_or_add_inline().add_drawing('wp:inline')
        image_id._inline.graphic.graphicData.append(parse_xml(f'''
            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                    <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                        <pic:blipFill>
                            <a:blip r:embed="{image._r.get_or_add_drawing().get_or_add_inline().get_or_add_graphic().get_or_add_graphicData().get_or_add_pic().get_or_add_blipFill().blip.get('r:embed')}"/>
                            <a:stretch>
                                <a:fillRect/>
                            </a:stretch>
                        </pic:blipFill>
                        <pic:spPr>
                            <a:xfrm>
                                <a:ext cx="{image_width}" cy="{image_height}"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect">
                                <a:avLst/>
                            </a:prstGeom>
                        </pic:spPr>
                    </pic:pic>
                </a:graphicData>
            </a:graphic>
        '''))

        output_paragraph = output_document.add_paragraph()
        output_run = output_paragraph.add_run()
        output_run._r.append(image_id)
    
    return output_document

try:
    start_time = time.time()

    original_paragraphs, original_images, original_tables = extract_paragraphs(args.original_documentFile)
    summaries = []
    for paragraph in original_paragraphs:
        if paragraph.strip() != "":
            summary = summarize_paragraph(paragraph)
            summaries.append(summary)

    output_document = Document()

    for i, (paragraph, summary) in enumerate(zip(original_paragraphs, summaries)):
        output_document.add_paragraph(summary)

    copy_tables(original_tables, output_document)
    copy_images(original_images, output_document)

    
    output_document.save(args.new_documentFile)

    end_time = time.time()

    print(f"New file generated: {args.new_documentFile}, Time Duration: {(end_time - start_time)/60:.2f} minutes")

except Exception as error:
    print("Exception:", str(error))
